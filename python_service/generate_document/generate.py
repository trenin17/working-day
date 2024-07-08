import xml.etree.ElementTree as ET
from docx import Document
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from datetime import datetime
import subprocess
from aiohttp import web
from sign_document.stamp import create_stamp, StampData
from s3_client.aws_utils import upload_and_presign
from threading import Lock

mutex = Lock()

def convert_docx_to_pdf(docx_path, output_pdf_path):
    with mutex:
        cmd = [
            'libreoffice', '--headless', '--convert-to', 'pdf', '--outdir',
            output_pdf_path, docx_path
        ]
        subprocess.run(cmd, check=True)


def plural_form(number, first, second, third):
    one_digit = number % 10
    two_digit = number % 100

    if two_digit >= 10 and two_digit <= 20:
        return third

    if one_digit == 1:
        return first

    if one_digit >= 2 and one_digit <= 4:
        return second

    return third

def get_duration(start_date, end_date):
    start_dd = datetime.strptime(start_date, '%d.%m.%Y')
    end_dd = datetime.strptime(end_date, '%d.%m.%Y')

    # вычисляем разницу между датами
    difference = end_dd - start_dd

    # количество дней - это total_seconds разделить на количество секунд в дне (86400)
    duration = int(difference.total_seconds() / 86400) + 1
    duration = str(duration) + ' ' + plural_form(duration, 'календарный день', 'календарных дня', ' календарных дней')

    return duration

def replace_macros_in_word(doc_path, replacements, output_path):
    doc = Document(doc_path)

    # Замена в абзацах
    for paragraph in doc.paragraphs:
        for macro, value in replacements.items():
            if macro in paragraph.text:
                paragraph.text = paragraph.text.replace(macro, value)

    # Замена в таблицах
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for macro, value in replacements.items():
                    if macro in cell.text:
                        cell.text = cell.text.replace(macro, value)

    # Замена в текстовых блоках
    for element in doc.element.body:
        if element.tag == '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}p':
            # Это абзац
            text_elements = element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t')
            for t in text_elements:
                for macro, value in replacements.items():
                    if macro in t.text:
                        t.text = t.text.replace(macro, value)

    doc.save(output_path)

def delete_tmp_files(file_key):
    subprocess.run(['rm', '-rf', '/tmp/' + file_key + '*'])


async def generate_document(request):
    try:
        # Retrieve `file_key` from the query string
        file_key = request.rel_url.query['file_key']

        # Get JSON data from the body
        data = await request.json()
        request_type = data['request_type']

        employee_id = data['employee_id']
        employee_name = data['employee_name']
        employee_surname = data['employee_surname']
        employee_patronymic = data.get('employee_patronymic', "")
        employee_position = data.get('employee_position', "")

        company_name = data.get('subcompany', "")
        company_id = data.get('company_id', "")

        head_name = data['head_name']
        head_surname = data['head_surname']
        head_patronymic = data.get('head_patronymic', "")
        head_position = data.get('head_position', "")
        start_date = data['start_date']
        end_date = data['end_date']
        duration = get_duration(start_date, end_date)

        first_start_date = data.get('first_start_date', "")
        second_start_date = data.get('second_start_date', "")
        first_end_date = data.get('first_end_date', "")
        second_end_date = data.get('second_end_date', "")
        first_duration = ""
        if first_start_date and first_end_date:
            first_duration = get_duration(first_start_date, first_end_date)
        second_duration = ""
        if second_start_date and second_end_date:
            second_duration = get_duration(second_start_date, second_end_date)

        employee_initials = f"{employee_surname} {employee_name[0]}."
        if employee_patronymic:
            employee_initials += f"{employee_patronymic[0]}."

        head_initials = f"{head_surname} {head_name[0]}."
        if head_patronymic:
            head_initials += f"{head_patronymic[0]}."

        now_date = datetime.today().strftime('%d.%m.%Y')

        replacements = {
            '%employee_name%': employee_name,
            '%employee_surname%': employee_surname,
            '%employee_patronymic%': employee_patronymic,
            '%employee_position%': employee_position,
            '%employee_initials%': employee_initials,
            '%company_name%': company_name,
            '%head_name%': head_name,
            '%head_surname%': head_surname,
            '%head_patronymic%': head_patronymic,
            '%head_position%': head_position,
            '%head_initials%': head_initials,
            '%start_date%': start_date,
            '%end_date%': end_date,
            '%duration%': duration,
            '%now_date%': now_date,
            '%first_start_date%': first_start_date,
            '%first_end_date%': first_end_date,
            '%first_duration%': first_duration,
            '%second_start_date%': second_start_date,
            '%second_end_date%': second_end_date,
            '%second_duration%': second_duration
        }

        file_name = file_key + '.docx'
        output_path_word = '/tmp/' + file_name
        replace_macros_in_word("generate_document/templates/" + company_id + "_" + request_type + ".docx",
                                replacements, output_path_word)

        output_path_pdf = '/tmp/' + file_key + '.pdf'

        convert_docx_to_pdf(output_path_word, '/tmp')

        output_path_pdf_signed = '/tmp/' + file_key + '_signed.pdf'
        # TODO: change company name
        stamp_data = StampData(now_date, employee_initials, company_name, employee_id, file_key)
        create_stamp(output_path_pdf, output_path_pdf_signed, stamp_data)

        url = upload_and_presign(output_path_pdf_signed, file_key + '.pdf')
        delete_tmp_files(file_key)
        return web.Response(status=200, content_type='text/plain', text=url)

    except Exception as e:
        return web.Response(status=500, text=str(e))
