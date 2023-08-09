import xml.etree.ElementTree as ET
from docx import Document
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import boto3
import json
import base64
from datetime import datetime

storage_client = None

def get_boto_session():
    boto_session = boto3.session.Session(
        aws_access_key_id='YCAJESGfHPZvOGYgqBjDkrLCZ',
        aws_secret_access_key='YCOFwp-4AxDaEEKeDBPx6YVuvU8Bqx-Z3hcQdIR8'
    )
    return boto_session

def get_storage_client():
    global storage_client
    if storage_client is not None:
        return storage_client

    storage_client = get_boto_session().client(
        service_name='s3',
        endpoint_url='https://storage.yandexcloud.net',
        region_name='ru-central1'
    )
    return storage_client

def upload_and_presign(file_path, object_name):
    client = get_storage_client()
    bucket = 'trenin17-results'
    client.upload_file(file_path, bucket, object_name)
    return client.generate_presigned_url('get_object', Params={'Bucket': bucket, 'Key': object_name}, ExpiresIn=3600)

def plural_form(number, first, second, third):
    one_digit = number % 10
    two_digit = number % 100

    if two_digit >= 10 and two_digit <= 20:
        return third
    
    if one_digit == 1:
        return first

    if one_digit >= 2 and one_digit <= 4:
        return second

    return third
    
def duration(start_date, end_date):
    start_dd = datetime.strptime(start_date, '%d.%m.%Y')
    end_dd = datetime.strptime(end_date, '%d.%m.%Y')

    # вычисляем разницу между датами
    difference = end_dd - start_dd

    # количество дней - это total_seconds разделить на количество секунд в дне (86400)
    duration = int(difference.total_seconds() / 86400) + 1
    duration = str(duration) + ' ' + plural_form(duration, 'календарный день', 'календарных дня', ' календарных дней')

    return duration

def replace_macros_in_word(doc_path, replacements, output_path):
    doc = Document(doc_path)

    # Замена в абзацах
    for paragraph in doc.paragraphs:
        for macro, value in replacements.items():
            if macro in paragraph.text:
                paragraph.text = paragraph.text.replace(macro, value)

    # Замена в таблицах
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for macro, value in replacements.items():
                    if macro in cell.text:
                        cell.text = cell.text.replace(macro, value)

    # Замена в текстовых блоках
    for element in doc.element.body:
        if element.tag == '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}p':
            # Это абзац
            text_elements = element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t')
            for t in text_elements:
                for macro, value in replacements.items():
                    if macro in t.text:
                        t.text = t.text.replace(macro, value)

    doc.save(output_path)

def handler(event, context):
    file_key = event['queryStringParameters']['file_key']
    body = json.loads(base64.b64decode(event['body']))
    request_type = body['request_type']

    employee_name = body['employee_name']
    employee_surname = body['employee_surname']
    employee_patronymic = body.get('employee_patronymic', "")
    employee_position = body.get('employee_position', "")
    
    head_name = body['head_name']
    head_surname = body['head_surname']
    head_patronymic = body.get('head_patronymic', "")
    head_position = body.get('head_position', "")
    start_date = body['start_date']
    end_date = body['end_date']
    duration = duration(start_date, end_date)

    first_start_date = body.get('first_start_date', "")
    second_start_date = body.get('second_start_date', "")
    first_end_date = body.get('first_end_date', "")
    second_end_date = body.get('second_end_date', "")
    first_duration = duration(first_start_date, first_end_date)
    second_duration = duration(second_start_date, second_end_date)

    employee_initials = ""
    if employee_patronymic == "":
        employee_initials = employee_surname + " " + employee_name[0] + "."
    else:
        employee_initials = employee_surname + " " + employee_name[0] + "." + employee_patronymic[0] + "."

    head_initials = ""
    if head_patronymic == "":
        head_initials = head_surname + " " + head_name[0] + "."
    else:
        head_initials = head_surname + " " + head_name[0] + "." + head_patronymic[0] + "."

    now_date = datetime.today().strftime('%d.%m.%Y')

    replacements = {
        '%employee_name%': employee_name,
        '%employee_surname%': employee_surname,
        '%employee_patronymic%': employee_patronymic,
        '%employee_position%': employee_position,
        '%employee_initials%': employee_initials,
        '%head_name%': head_name,
        '%head_surname%': head_surname,
        '%head_patronymic%': head_patronymic,
        '%head_position%': head_position,
        '%head_initials%': head_initials,
        '%start_date%': start_date,
        '%end_date%': end_date,
        '%duration%': duration,
        '%now_date%': now_date,
        '%first_start_date%': first_start_date,
        '%first_end_date%': first_end_date,
        '%first_duration%': first_duration,
        '%second_start_date%': second_start_date,
        '%second_end_date%': second_end_date,
        '%second_duration%': second_duration
    }

    file_key += '.docx'
    output_path = '/tmp/' + file_key
    replace_macros_in_word("templates/" + request_type + ".docx", replacements, output_path)

    url = upload_and_presign(output_path, file_key)
    return {
        'statusCode': 200,
        'headers': {
            'Content-Type': 'text/plain'
        },
        'isBase64Encoded': False,
        'body': url
    }
